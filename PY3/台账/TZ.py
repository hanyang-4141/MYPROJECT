from docx import Document
from docxcompose.composer import Composer

import pandas as PD
doc_path = r'NULLTABLE.docx'
excel_path = r'燃料2.0台账表格.xlsx'

df_A1 = PD.read_excel(excel_path, sheet_name='A1设备技术信息记录表')
df_A1 = df_A1.fillna('无')
df_A2 = PD.read_excel(excel_path, sheet_name='A2附属设备清单记录表')
df_A2 = df_A2.fillna('无')
df_A3 = PD.read_excel(excel_path, sheet_name='A3备品配件清单记录表')
df_A3 = df_A3.fillna('无')
df_B1 = PD.read_excel(excel_path, sheet_name='B1设备检修记录表')
df_B1 = df_B1.fillna('无')
df_C1 = PD.read_excel(excel_path, sheet_name='C1设备日常维护记录表')
df_C1 = df_C1.fillna('无')
df_D1 = PD.read_excel(excel_path, sheet_name='D1设备事故事件记录表')
df_D1 = df_D1.fillna('无')
df_E1 = PD.read_excel(excel_path, sheet_name='E1设备更新改造记录表')
df_E1 = df_E1.fillna('无')





def insert_data(table, cell_i, cell_j, data):
    # print(data)
    if data == '无':
        # print('nan')
        return
    else:
        table.cell(cell_i, cell_j).text = data


for i in range(df_A1.shape[0]):
    print(i)
    MyDoc = Document(doc_path)
    for paragraph in MyDoc.paragraphs:
        # print(paragraph.text)
        line_text = paragraph.text
        if line_text.find("此处为设备名称") != -1:
            paragraph.text = df_A1.at[i, '设备名称']
    tables = MyDoc.tables
    table_A1 = tables[0]
    table_A2 = tables[1]
    table_A3 = tables[2]
    table_B1 = tables[3]
    table_C1 = tables[4]
    table_D1 = tables[5]
    table_E1 = tables[6]
    # print(df_A1.shape[0])
    # mydocx.aMyDoc_heading(df_A1.at[i, '设备名称'], level=4)
    # mydocx.aMyDoc_heading('   基础信息', level=4)
    # mydocx.aMyDoc_heading('A1  设备技术信息记录表', level=5)
    #**************************************************************************A1*********
    insert_data(table_A1, 0, 1, df_A1.at[i, '设备名称'])
    insert_data(table_A1, 0, 3, df_A1.at[i, '设备编码'])
    insert_data(table_A1, 1, 1, df_A1.at[i, '设备型号'])
    insert_data(table_A1, 1, 3, df_A1.at[i, '制造厂家'])
    insert_data(table_A1, 2, 1, df_A1.at[i, '设备类型'])
    insert_data(table_A1, 2, 3, df_A1.at[i, '安装位置'])
    insert_data(table_A1, 3, 1, df_A1.at[i, '投运日期'])
    insert_data(table_A1, 3, 3, df_A1.at[i, '所属部门'])
    insert_data(table_A1, 4, 1, df_A1.at[i, '所属专业'])
    insert_data(table_A1, 4, 3, df_A1.at[i, '责任人'])
    insert_data(table_A1, 6, 1, df_A1.at[i, '技术参数'])

    # **************************************************************************A2*********
    A2附属 = df_A2.at[i, '附属']
    if A2附属 != '无':
        附属list = A2附属.splitlines()
        print(附属list)
        for index, 附属 in enumerate(附属list):
            TempList = 附属.split('，')
            print(index)
            print(TempList)
            insert_data(table_A2, index + 1, 0, TempList[0])
            insert_data(table_A2, index + 1, 1, TempList[1])
            insert_data(table_A2, index + 1, 2, TempList[2])
            insert_data(table_A2, index + 1, 3, TempList[3])
            insert_data(table_A2, index + 1, 4, TempList[4])
            insert_data(table_A2, index + 1, 5, TempList[5])
            insert_data(table_A2, index + 1, 6, TempList[6])

    # **************************************************************************A3*********
    A3备件 = df_A3.at[i, '备件清单']
    if A3备件 != '无':
        备件list = A3备件.splitlines()
        for index, 备件 in enumerate(备件list):
            TempList = 备件.split('，')
            insert_data(table_A3, index + 1, 0, TempList[0])
            insert_data(table_A3, index + 1, 1, TempList[1])
            insert_data(table_A3, index + 1, 2, TempList[2])
            insert_data(table_A3, index + 1, 3, TempList[3])
            insert_data(table_A3, index + 1, 4, TempList[4])
            insert_data(table_A3, index + 1, 5, TempList[5])
            insert_data(table_A3, index + 1, 6, TempList[6])


    # **************************************************************************B1*********
    insert_data(table_B1, 0, 1, df_B1.at[i, '设备名称'])
    insert_data(table_B1, 0, 3, df_B1.at[i, '检修性质'])
    insert_data(table_B1, 1, 1, df_B1.at[i, '检修开始时间'])
    insert_data(table_B1, 1, 3, df_B1.at[i, '检修结束时间'])
    insert_data(table_B1, 2, 1, df_B1.at[i, '修前状况'])
    insert_data(table_B1, 3, 1, df_B1.at[i, '检修主要内容'])
    insert_data(table_B1, 4, 1, df_B1.at[i, '试运情况'])
    insert_data(table_B1, 5, 3, df_B1.at[i, '修后状况'])
    insert_data(table_B1, 6, 1, df_B1.at[i, '备品配件'])
    insert_data(table_B1, 7, 3, df_B1.at[i, '检修人员'])

    
    
    # **************************************************************************C1*********
    insert_data(table_C1, 0, 1, df_C1.at[i, '设备名称'])
    insert_data(table_C1, 0, 3, df_C1.at[i, '维护时间'])
    insert_data(table_C1, 1, 1, df_C1.at[i, '维护性质'])
    insert_data(table_C1, 1, 3, df_C1.at[i, '维护人员'])
    insert_data(table_C1, 2, 1, df_C1.at[i, '维护内容'])
    insert_data(table_C1, 3, 1, df_C1.at[i, '维护备件'])

    # **************************************************************************D1*********
    insert_data(table_D1, 0, 1, df_D1.at[i, '设备名称'])
    insert_data(table_D1, 1, 1, df_D1.at[i, '事件名称'])
    insert_data(table_D1, 2, 1, df_D1.at[i, '事件时间'])
    insert_data(table_D1, 2, 3, df_D1.at[i, '事件性质'])
    insert_data(table_D1, 4, 1, df_D1.at[i, '事件简述'])
    insert_data(table_D1, 5, 1, df_D1.at[i, '原因分析'])
    insert_data(table_D1, 6, 1, df_D1.at[i, '处理过程'])
    insert_data(table_D1, 7, 1, df_D1.at[i, '防范措施'])
    insert_data(table_D1, 8, 1, df_D1.at[i, '处理人员'])
    insert_data(table_D1, 8, 3, df_D1.at[i, '消除时间'])
    # **************************************************************************E1*********
    insert_data(table_E1, 0, 1, df_E1.at[i, '设备名称'])
    insert_data(table_E1, 0, 3, df_E1.at[i, '更改性质'])
    insert_data(table_E1, 1, 1, df_E1.at[i, '起止时间'])
    insert_data(table_E1, 1, 3, df_E1.at[i, '更改投资'])
    insert_data(table_E1, 2, 1, df_E1.at[i, '参与人员'])
    insert_data(table_E1, 3, 1, df_E1.at[i, '更改原因'])
    insert_data(table_E1, 4, 1, df_E1.at[i, '更改内容'])
    insert_data(table_E1, 5, 1, df_E1.at[i, '主要供货厂家及施工单位'])
    insert_data(table_E1, 6, 1, df_E1.at[i, '更改效果'])
    insert_data(table_E1, 7, 1, df_E1.at[i, '更改技术资料'])



    MyDoc.save('1\hehe{}.docx'.format(i))
    MyDoc = None




